from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io, os, datetime, sqlite3, smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from pdf_builder import build_pdf

app = Flask(__name__)

OWNER_NAME = "יהודה קורץ"
OWNER_TAX  = "027394865"
DB_PATH    = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data.db")
def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db()
    conn.execute('''CREATE TABLE IF NOT EXISTS clients (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        tax_id TEXT,
        address TEXT,
        phone TEXT
    )''')
    conn.execute('''CREATE TABLE IF NOT EXISTS documents (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        doc_type TEXT NOT NULL,
        doc_num INTEGER NOT NULL,
        date TEXT NOT NULL,
        client_name TEXT NOT NULL,
        client_id TEXT,
        address TEXT,
        description TEXT,
        inv_amt REAL DEFAULT 0,
        bank_amt REAL DEFAULT 0,
        nik_amt REAL DEFAULT 0,
        created_at TEXT DEFAULT (datetime('now','localtime'))
    )''')
    conn.execute('''CREATE TABLE IF NOT EXISTS counters (
        doc_type TEXT PRIMARY KEY,
        next_num INTEGER NOT NULL
    )''')
    for t, n in [("חשבונית עסקה", 2600166), ("קבלה", 2601166), ("חשבונית עסקה + קבלה", 26002166)]:
        conn.execute("INSERT OR IGNORE INTO counters VALUES (?, ?)", (t, n))
    conn.commit()
    conn.close()

init_db()

def next_doc_num(conn, doc_type):
    row = conn.execute("SELECT next_num FROM counters WHERE doc_type=?", (doc_type,)).fetchone()
    num = row["next_num"]
    conn.execute("UPDATE counters SET next_num=? WHERE doc_type=?", (num+1, doc_type))
    return num
    
    def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_rtl(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)

def make_border(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_para(cell, text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.RIGHT):
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    para.alignment = align
    set_rtl(para)
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return para
    def build_doc(data):
    doc_type  = data.get('doc_type', '')
    doc_num   = data.get('doc_num', '')
    date_str  = data.get('date', datetime.date.today().strftime('%d/%m/%Y'))
    client    = data.get('client', '')
    client_id = data.get('client_id', '')
    address   = data.get('address', '')
    desc      = data.get('desc', '')
    inv_amt   = data.get('inv_amt')
    bank_amt  = data.get('bank_amt')
    nik_amt   = data.get('nik_amt')
    pct_nik   = data.get('pct_nik', 0.2)

    def sf(v):
        try: return float(v) if v else 0.0
        except: return 0.0

    inv_amt = sf(inv_amt)
    bank_amt = sf(bank_amt)
    nik_amt  = sf(nik_amt)
    pct_nik  = sf(pct_nik) or 0.2
    rec_total = bank_amt + nik_amt

    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2)
    section.top_margin  = section.bottom_margin = Cm(2)
    doc.settings.element.append(OxmlElement('w:bidi'))

    for style_name in ['Normal', 'Table Grid']:
        try:
            s = doc.styles[style_name]
            pPr = s.element.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
        except: pass

    title_color = {'חשבונית עסקה':'2563EB','קבלה':'16A34A'}.get(doc_type,'7C3AED')

    for txt, sz, bold, col in [
        (doc_type, 26, True, title_color),
        (OWNER_NAME, 18, True, '1E3A5F'),
        (f"עוסק פטור מס׳ {OWNER_TAX}", 12, False, '374151'),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_rtl(p)
        r = p.add_run(txt)
        r.font.name='Arial'; r.font.size=Pt(sz); r.font.bold=bold
        r.font.color.rgb = RGBColor.from_string(col)

    doc.add_paragraph()
    t1 = doc.add_table(rows=1, cols=2)
    t1.style = 'Table Grid'

    c_date = t1.rows[0].cells[0]
    set_cell_bg(c_date, "F8FAFC")
    p_date = c_date.paragraphs[0]
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p_date)
    rd = p_date.add_run(f"תאריך: {date_str}")
    rd.font.name='Arial'; rd.font.size=Pt(12)

    c_num = t1.rows[0].cells[1]
    set_cell_bg(c_num, "EFF6FF")
    make_border(c_num, "2563EB")
    p_num = c_num.paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p_num)
    rl = p_num.add_run("מספר מסמך:  ")
    rl.font.name='Arial'; rl.font.size=Pt(12); rl.font.bold=True
    rl.font.color.rgb = RGBColor.from_string('2563EB')
    rn = p_num.add_run(str(doc_num))
    rn.font.name='Arial'; rn.font.size=Pt(16); rn.font.bold=True
    rn.font.color.rgb = RGBColor.from_string('1E3A5F')

    doc.add_paragraph()

    t2 = doc.add_table(rows=3, cols=2)
    t2.style = 'Table Grid'

    c_title = t2.rows[0].cells[0]
    c_title.merge(t2.rows[0].cells[1])
    set_cell_bg(c_title, "E8F5E9")
    p_title = c_title.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p_title)
    rt = p_title.add_run("פרטי הלקוח")
    rt.font.name='Arial'; rt.font.size=Pt(12); rt.font.bold=True
    rt.font.color.rgb = RGBColor.from_string('16A34A')

    c_name = t2.rows[1].cells[1]
    set_cell_bg(c_name, "F8FAFC")
    p_name = c_name.paragraphs[0]; p_name.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_rtl(p_name)
    r_name = p_name.add_run(f"שם לקוח: {client}")
    r_name.font.name='Arial'; r_name.font.size=Pt(11)

    c_id = t2.rows[1].cells[0]
    set_cell_bg(c_id, "F8FAFC")
    p_id = c_id.paragraphs[0]; p_id.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_rtl(p_id)
    r_id = p_id.add_run(f"מס׳ עוסק: {client_id}")
    r_id.font.name='Arial'; r_id.font.size=Pt(11)

    c_addr = t2.rows[2].cells[1]
    set_cell_bg(c_addr, "F8FAFC")
    p_addr = c_addr.paragraphs[0]; p_addr.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_rtl(p_addr)
    r_addr = p_addr.add_run(f"כתובת: {address}")
    r_addr.font.name='Arial'; r_addr.font.size=Pt(11)

    c_desc = t2.rows[2].cells[0]
    set_cell_bg(c_desc, "F8FAFC")
    p_desc = c_desc.paragraphs[0]; p_desc.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_rtl(p_desc)
    r_desc = p_desc.add_run("")
    r_desc.font.name='Arial'; r_desc.font.size=Pt(11)

    doc.add_paragraph()
    if doc_type in ('חשבונית עסקה','חשבונית עסקה + קבלה'):
        t3 = doc.add_table(rows=2, cols=3)
        t3.style = 'Table Grid'
        headers = ["סכום ₪", "כמות", "תיאור השירות / הפריט"]
        values  = [f"{inv_amt:,.2f}", "1", desc or ""]
        for i, h in enumerate(headers):
            c = t3.rows[0].cells[i]
            set_cell_bg(c, "2563EB")
            p = c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_rtl(p)
            r = p.add_run(h)
            r.font.name='Arial'; r.font.size=Pt(11); r.font.bold=True
            r.font.color.rgb = RGBColor.from_string('FFFFFF')
        for i, val in enumerate(values):
            c = t3.rows[1].cells[i]
            set_cell_bg(c, "EFF6FF")
            p = c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_rtl(p)
            r = p.add_run(val)
            r.font.name='Arial'; r.font.size=Pt(11)

        doc.add_paragraph()

        t4 = doc.add_table(rows=1, cols=2)
        t4.style = 'Table Grid'
        cr = t4.rows[0].cells[0]
        cl = t4.rows[0].cells[1]
        set_cell_bg(cr, "2563EB")
        pr = cr.paragraphs[0]; pr.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_rtl(pr)
        rr = pr.add_run('סה"כ לתשלום:')
        rr.font.name='Arial'; rr.font.size=Pt(13); rr.font.bold=True
        rr.font.color.rgb = RGBColor.from_string('FFFFFF')
        set_cell_bg(cl, "DBEAFE")
        pl = cl.paragraphs[0]; pl.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_rtl(pl)
        rl2 = pl.add_run(f"₪ {inv_amt:,.2f}")
        rl2.font.name='Arial'; rl2.font.size=Pt(14); rl2.font.bold=True
        rl2.font.color.rgb = RGBColor.from_string('1E3A5F')
        doc.add_paragraph()

    if doc_type in ('קבלה','חשבונית עסקה + קבלה'):
        rows_data = [
            ('סה"כ שולם:', f"₪ {rec_total:,.2f}", "F0FDF4","16A34A"),
            ("התקבל בבנק:", f"₪ {bank_amt:,.2f}", "F0FDF4","15803D"),
            (f"ניכוי במקור ({int(pct_nik*100)}%):", f"₪ {nik_amt:,.2f}", "F5F3FF","7C3AED"),
            ("אמצעי תשלום:", "____________", "F8FAFC","374151"),
        ]
        t5 = doc.add_table(rows=4, cols=2)
        t5.style = 'Table Grid'
        for i,(label,val,bg,col) in enumerate(rows_data):
            c_label = t5.rows[i].cells[1]
            c_val   = t5.rows[i].cells[0]
            set_cell_bg(c_label, bg); set_cell_bg(c_val, bg)
            p_l = c_label.paragraphs[0]; p_l.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_rtl(p_l)
            rl = p_l.add_run(label)
            rl.font.name='Arial'; rl.font.size=Pt(11); rl.font.bold=True
            rl.font.color.rgb = RGBColor.from_string(col)
            p_v = c_val.paragraphs[0]; p_v.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_rtl(p_v)
            rv = p_v.add_run(val)
            rv.font.name='Arial'; rv.font.size=Pt(12); rv.font.bold=True
            rv.font.color.rgb = RGBColor.from_string(col)
        doc.add_paragraph()

    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p_note)
    rn = p_note.add_run('* עוסק פטור — אינו מחייב במע"מ')
    rn.font.name='Arial'; rn.font.size=Pt(9); rn.font.italic=True
    rn.font.color.rgb = RGBColor.from_string('6B7280')

    doc.add_paragraph()

    t6 = doc.add_table(rows=1, cols=2)
    t6.style = 'Table Grid'
    add_para(t6.rows[0].cells[1], "חתימת המוכר: ____________________", size=11)
    add_para(t6.rows[0].cells[0], "חתימת הלקוח: ____________________", size=11)

    return doc
    @app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate():
    try:
        data = request.json
        doc_type = data.get('doc_type','')
        if doc_type not in ["חשבונית עסקה","קבלה","חשבונית עסקה + קבלה"]:
            return jsonify({'error':'סוג מסמך לא חוקי'}), 400

        def sf(v):
            try: return float(v) if v else 0.0
            except: return 0.0

        inv_amt  = sf(data.get('inv_amt'))
        bank_amt = sf(data.get('bank_amt'))
        nik_amt  = sf(data.get('nik_amt'))

        conn = get_db()
        doc_num = next_doc_num(conn, doc_type)
        data['doc_num'] = doc_num

        conn.execute('''INSERT INTO documents
            (doc_type,doc_num,date,client_name,client_id,address,description,inv_amt,bank_amt,nik_amt)
            VALUES (?,?,?,?,?,?,?,?,?,?)''', (
            doc_type, doc_num,
            data.get('date',''),
            data.get('client',''),
            data.get('client_id',''),
            data.get('address',''),
            data.get('desc',''),
            inv_amt, bank_amt, nik_amt,
        ))
        conn.commit()
        conn.close()

        doc = build_doc(data)
        buf = io.BytesIO()
        doc.save(buf); buf.seek(0)

        safe_client = (data.get('client','') or '').replace(' ','_')
        safe_type   = doc_type.replace(' ','_').replace('+','-')
        filename    = f"{safe_type}_{doc_num}_{safe_client}.docx"

        return send_file(buf, as_attachment=True, download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/generate_pdf', methods=['POST'])
def generate_pdf():
    try:
        data = request.json
        doc_type = data.get('doc_type','')
        if doc_type not in ["חשבונית עסקה","קבלה","חשבונית עסקה + קבלה"]:
            return jsonify({'error':'סוג מסמך לא חוקי'}), 400

        def sf(v):
            try: return float(v) if v else 0.0
            except: return 0.0

        inv_amt  = sf(data.get('inv_amt'))
        bank_amt = sf(data.get('bank_amt'))
        nik_amt  = sf(data.get('nik_amt'))

        conn = get_db()
        doc_num = next_doc_num(conn, doc_type)
        data['doc_num'] = doc_num

        conn.execute('''INSERT INTO documents
            (doc_type,doc_num,date,client_name,client_id,address,description,inv_amt,bank_amt,nik_amt)
            VALUES (?,?,?,?,?,?,?,?,?,?)''', (
            doc_type, doc_num,
            data.get('date',''),
            data.get('client',''),
            data.get('client_id',''),
            data.get('address',''),
            data.get('desc',''),
            inv_amt, bank_amt, nik_amt,
        ))
        conn.commit()
        conn.close()

        buf = build_pdf(data)
        safe_client = (data.get('client','') or '').replace(' ','_')
        safe_type   = doc_type.replace(' ','_').replace('+','-')
        filename    = f"{safe_type}_{doc_num}_{safe_client}.pdf"

        return send_file(buf, as_attachment=True, download_name=filename,
            mimetype='application/pdf')

    except Exception as e:
        return jsonify({'error': str(e)}), 500
        @app.route('/clients', methods=['GET','POST'])
def clients():
    conn = get_db()
    if request.method == 'GET':
        rows = conn.execute("SELECT * FROM clients ORDER BY name").fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    else:
        data = request.json
        conn.execute("INSERT INTO clients (name,tax_id,address,phone) VALUES (?,?,?,?)",
            (data.get('name',''), data.get('id',''), data.get('address',''), data.get('phone','')))
        conn.commit(); conn.close()
        return jsonify({'ok':True})

@app.route('/history')
def history():
    client_name = request.args.get('client','')
    conn = get_db()
    if client_name:
        rows = conn.execute(
            "SELECT * FROM documents WHERE client_name=? ORDER BY id DESC", (client_name,)
        ).fetchall()
    else:
        rows = conn.execute("SELECT * FROM documents ORDER BY id DESC LIMIT 200").fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/ledger/<client_name>')
def ledger(client_name):
    conn = get_db()
    rows = conn.execute(
        "SELECT * FROM documents WHERE client_name=? ORDER BY id ASC", (client_name,)
    ).fetchall()
    conn.close()
    docs = [dict(r) for r in rows]

    total_debit  = 0.0
    total_bank   = 0.0
    total_nik    = 0.0
    total_credit = 0.0
    entries = []
    balance = 0.0

    for d in docs:
        debit  = 0.0
        credit = 0.0
        if d['doc_type'] in ('חשבונית עסקה', 'חשבונית עסקה + קבלה'):
            debit = d['inv_amt'] or 0.0
            total_debit += debit
        if d['doc_type'] in ('קבלה', 'חשבונית עסקה + קבלה'):
            bank = d['bank_amt'] or 0.0
            nik  = d['nik_amt']  or 0.0
            credit = bank + nik
            total_bank += bank
            total_nik  += nik
            total_credit += credit
        balance += debit - credit
        entries.append({**d, 'debit': debit, 'credit': credit, 'balance': balance})

    return jsonify({
        'client': client_name, 'entries': entries,
        'total_debit': total_debit, 'total_bank': total_bank,
        'total_nik': total_nik, 'total_credit': total_credit, 'balance': balance,
    })

@app.route('/client_names')
def client_names():
    conn = get_db()
    rows = conn.execute(
        "SELECT DISTINCT client_name FROM documents ORDER BY client_name"
    ).fetchall()
    conn.close()
    return jsonify([r['client_name'] for r in rows])

@app.route('/document/<int:doc_id>', methods=['GET','PUT','DELETE'])
def document(doc_id):
    conn = get_db()
    if request.method == 'GET':
        row = conn.execute("SELECT * FROM documents WHERE id=?", (doc_id,)).fetchone()
        conn.close()
        if not row: return jsonify({'error':'לא נמצא'}), 404
        return jsonify(dict(row))
    elif request.method == 'PUT':
        data = request.json
        def sf(v):
            try: return float(v) if v else 0.0
            except: return 0.0
        conn.execute('''UPDATE documents SET
            date=?, client_name=?, client_id=?, address=?, description=?,
            inv_amt=?, bank_amt=?, nik_amt=?
            WHERE id=?''', (
            data.get('date',''), data.get('client',''), data.get('client_id',''),
            data.get('address',''), data.get('desc',''),
            sf(data.get('inv_amt')), sf(data.get('bank_amt')), sf(data.get('nik_amt')),
            doc_id
        ))
        conn.commit(); conn.close()
        return jsonify({'ok': True})
    elif request.method == 'DELETE':
        conn.execute("DELETE FROM documents WHERE id=?", (doc_id,))
        conn.commit(); conn.close()
        return jsonify({'ok': True})

@app.route('/send_email', methods=['POST'])
def send_email():
    try:
        data = request.json
        to_email  = data.get('to_email','').strip()
        doc_type  = data.get('doc_type','')
        doc_num   = data.get('doc_num','')
        client    = data.get('client','')
        fmt       = data.get('format','pdf')

        if not to_email:
            return jsonify({'error': 'נא להזין כתובת מייל'}), 400

        gmail_user = os.environ.get('GMAIL_USER','')
        gmail_pass = os.environ.get('GMAIL_PASS','')
        if not gmail_user or not gmail_pass:
            return jsonify({'error': 'הגדרות מייל חסרות בשרת'}), 500

        if fmt == 'pdf':
            file_buf  = build_pdf(data)
            filename  = f"{doc_type}_{doc_num}_{client}.pdf"
            mime_type = 'application/pdf'
        else:
            doc = build_doc(data)
            file_buf = io.BytesIO()
            doc.save(file_buf); file_buf.seek(0)
            filename  = f"{doc_type}_{doc_num}_{client}.docx"
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        msg = MIMEMultipart()
        msg['From']    = gmail_user
        msg['To']      = to_email
        msg['Subject'] = f'{doc_type} מספר {doc_num} — יהודה קורץ'

        body = f"שלום,\n\nמצורף {doc_type} מספר {doc_num}.\n\nבברכה,\nיהודה קורץ\nעוסק פטור מס׳ 027394865"
        msg.attach(MIMEText(body, 'plain', 'utf-8'))

        part = MIMEBase('application', 'octet-stream')
        part.set_payload(file_buf.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename="{filename}"')
        part.add_header('Content-Type', mime_type)
        msg.attach(part)

        with smtplib.SMTP('smtp.gmail.com', 587) as server:
            server.ehlo()
            server.starttls()
            server.login(gmail_user, gmail_pass)
            server.sendmail(gmail_user, to_email, msg.as_string())

        return jsonify({'ok': True, 'message': f'המייל נשלח ל-{to_email}'})

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/reset_data', methods=['POST'])
def reset_data():
    conn = get_db()
    conn.execute("DELETE FROM documents")
    conn.execute("UPDATE counters SET next_num=2600166 WHERE doc_type='חשבונית עסקה'")
    conn.execute("UPDATE counters SET next_num=2601166 WHERE doc_type='קבלה'")
    conn.execute("UPDATE counters SET next_num=26002166 WHERE doc_type='חשבונית עסקה + קבלה'")
    conn.commit()
    conn.close()
    return jsonify({'ok': True, 'message': 'כל המסמכים אופסו — הלקוחות נשמרו'})

@app.route('/reset_all', methods=['POST'])
def reset_all():
    conn = get_db()
    conn.execute("DELETE FROM documents")
    conn.execute("DELETE FROM clients")
    conn.execute("UPDATE counters SET next_num=2600166 WHERE doc_type='חשבונית עסקה'")
    conn.execute("UPDATE counters SET next_num=2601166 WHERE doc_type='קבלה'")
    conn.execute("UPDATE counters SET next_num=26002166 WHERE doc_type='חשבונית עסקה + קבלה'")
    conn.commit()
    conn.close()
    return jsonify({'ok': True, 'message': 'כל הנתונים נמחקו כולל לקוחות'})

if __name__ == '__main__':
    app.run(debug=True)
